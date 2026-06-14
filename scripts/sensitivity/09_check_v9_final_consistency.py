from __future__ import annotations

import subprocess
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[2]
LOCAL_SITE_PACKAGES = PROJECT_ROOT / ".python_packages"
if LOCAL_SITE_PACKAGES.exists():
    sys.path.insert(0, str(LOCAL_SITE_PACKAGES))

from docx import Document
from PIL import Image
import pandas as pd

OUT_ROOT = PROJECT_ROOT / "output/submission_enhancement_v2"
TABLE_DIR = OUT_ROOT / "tables"
LOG_DIR = OUT_ROOT / "logs"
MANIFEST_DIR = OUT_ROOT / "manifests"
V8 = PROJECT_ROOT / "manuscript/BMC_Geriatrics_possible_sarcopenia_revised_v8_robustness_figures_integrated.docx"
V9 = PROJECT_ROOT / "manuscript/BMC_Geriatrics_possible_sarcopenia_revised_v9_final_consistency_corrected.docx"


def status_line(ok: bool, message: str, critical: bool = True) -> tuple[str, str, bool]:
    return ("PASS" if ok else ("FAIL" if critical else "WARNING"), message, critical)


def close(a: float, b: float, tol: float = 5e-4) -> bool:
    return abs(float(a) - float(b)) <= tol


def pdf_check(path: Path) -> tuple[bool, str]:
    if not path.exists():
        return False, "missing"
    b = path.read_bytes()
    ok = (
        b.startswith(b"%PDF")
        and b.rstrip().endswith(b"%%EOF")
        and b.count(b"/FontFile2") > 0
        and b.count(b"/Subtype /Type3") == 0
        and path.stat().st_size < 10 * 1024 * 1024
    )
    return ok, f"size_mb={path.stat().st_size/1024/1024:.3f}; FontFile2={b.count(b'/FontFile2')}; Type3={b.count(b'/Subtype /Type3')}"


def png_check(path: Path) -> tuple[bool, str]:
    if not path.exists():
        return False, "missing"
    im = Image.open(path)
    dpi = im.info.get("dpi", (0, 0))
    ok = min(dpi) >= 250 and path.stat().st_size < 10 * 1024 * 1024
    return ok, f"{im.size[0]}x{im.size[1]}; dpi={dpi}; size_mb={path.stat().st_size/1024/1024:.3f}"


def render_v9() -> tuple[bool, str]:
    render_script = PROJECT_ROOT / "tools/render_docx.py"
    render_dir = LOG_DIR / "v9_rendered_pages"
    render_dir.mkdir(parents=True, exist_ok=True)
    if not render_script.exists():
        return False, f"render script not found: {render_script}"
    cmd = [sys.executable, str(render_script), str(V9), "--output_dir", str(render_dir), "--emit_pdf"]
    try:
        result = subprocess.run(
            cmd,
            cwd=str(PROJECT_ROOT),
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=180,
        )
    except Exception as exc:  # noqa: BLE001
        return False, f"render command failed to start: {exc}"
    pages = sorted(render_dir.glob("page-*.png"))
    if result.returncode == 0 and pages:
        return True, f"rendered {len(pages)} page PNGs in {render_dir.relative_to(PROJECT_ROOT)}"
    detail = (result.stderr or result.stdout or "no renderer output").strip().splitlines()
    return False, "; ".join(detail[-4:])


def main() -> None:
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    MANIFEST_DIR.mkdir(parents=True, exist_ok=True)
    checks: list[tuple[str, str, bool]] = []

    checks.append(status_line(V8.exists(), "v8 source manuscript exists"))
    checks.append(status_line(V9.exists(), "v9 manuscript created"))
    if not V9.exists():
        raise SystemExit("v9 manuscript is missing.")

    doc = Document(V9)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    lower = full_text.lower()

    stale_phrases = [
        "strict-label analysis was not available",
        "strict-label sensitivity analysis was not available",
        "strict-label analysis would strengthen future evaluation",
        "strict-label sensitivity analysis would strengthen future evaluation",
        "should be performed in future",
    ]
    for phrase in stale_phrases:
        checks.append(status_line(phrase not in lower, f"No stale phrase: {phrase}"))
    checks.append(status_line("following the outcome-component audit, we conducted a strict-label sensitivity analysis" in lower, "Strict-label Methods wording present"))
    checks.append(status_line("any observed positive component remained sufficient" in lower, "Strict-label positive rule described correctly"))
    checks.append(status_line("complete-component sample" not in lower, "Strict-label sample not described as complete-component sample"))

    # Table 2 values and format.
    expected_t2 = pd.read_csv(TABLE_DIR / "Table2_main_model_validation_concise.csv").astype(str)
    checks.append(status_line(len(doc.tables) == 3, "Manuscript has three editable Word tables"))
    t2 = doc.tables[1]
    checks.append(status_line(len(t2.rows) == 3 and len(t2.columns) == 8, "Table 2 is concise: 3 rows x 8 columns"))
    actual_t2 = [[cell.text.strip() for cell in row.cells] for row in t2.rows]
    expected_rows = [list(expected_t2.columns)] + expected_t2.values.tolist()
    checks.append(status_line(actual_t2 == expected_rows, "Table 2 values match authoritative concise validation table"))

    # Figure 2 common subset.
    fig2 = pd.read_csv(TABLE_DIR / "Figure2_calibration_summary.csv")
    basis = fig2[["n", "positive_n", "negative_n", "prevalence"]].drop_duplicates()
    checks.append(status_line(len(basis) == 1 and int(basis.iloc[0]["n"]) == 2430, "Figure 2 source uses common subset N=2430"))
    checks.append(status_line(int(basis.iloc[0]["positive_n"]) == 683 and int(basis.iloc[0]["negative_n"]) == 1747, "Figure 2 event/non-event counts match data packet"))
    checks.append(status_line("common paired subset of 2430 participants" in full_text, "Figure 2 common-subset basis stated in manuscript"))
    checks.append(status_line("differ slightly from the model-specific estimates reported in Table 2" in full_text, "Figure 2 vs Table 2 distinction stated"))

    # Figure 4.
    paired = pd.read_csv(TABLE_DIR / "paired_model_differences_with_95CI.csv")
    fig4 = pd.read_csv(TABLE_DIR / "Figure4_plot_data.csv")
    pmap = paired.set_index("metric")
    fmap = fig4[fig4["panel"].eq("A")].set_index("metric")
    for metric in ["AUROC advantage", "AUPRC advantage", "Brier advantage", "sensitivity difference", "specificity difference", "F1 difference"]:
        checks.append(status_line(metric in fmap.index and close(fmap.loc[metric, "estimate"], pmap.loc[metric, "estimate"]), f"Figure 4 Panel A {metric} estimate matches paired table"))
    fig4b = fig4[fig4["panel"].eq("B")]
    checks.append(status_line(set(fig4b["sample_basis"].dropna()) == {"common paired primary subset", "common paired strict-label subset"}, "Figure 4 Panel B declares common paired sample bases"))
    primary_b = fig4b[fig4b["label_definition"].eq("Primary available-component label")]
    strict_b = fig4b[fig4b["label_definition"].eq("Strict-label sensitivity")]
    checks.append(status_line((primary_b["n"] == 2430).all() and (primary_b["positive_n"] == 683).all(), "Figure 4 primary Panel B common N/events match packet"))
    checks.append(status_line((strict_b["n"] == 2373).all() and (strict_b["positive_n"] == 683).all(), "Figure 4 strict Panel B common N/events match packet"))
    checks.append(status_line("Brier-score advantage is defined as XGBoost Brier score minus Logistic Brier score" in full_text, "Figure 4 Brier-score direction stated"))
    checks.append(status_line("threshold-based differences use threshold 0.50" in full_text, "Figure 4 threshold metric basis stated"))

    # Strict label and thresholds.
    strict_flow = pd.read_csv(TABLE_DIR / "strict_label_flow_by_wave.csv")
    checks.append(status_line(int(strict_flow.loc[strict_flow["wave"].eq(1), "uncertain_partial_negative_n"].iloc[0]) == 313, "Strict Wave 1 uncertain count = 313"))
    checks.append(status_line(int(strict_flow.loc[strict_flow["wave"].eq(3), "uncertain_partial_negative_n"].iloc[0]) == 278, "Original Wave 3 uncertain count = 278"))
    checks.append(status_line("thresholds of 0.35 and 0.34" in full_text, "Screening thresholds 0.35 and 0.34 remain separate from default threshold"))
    checks.append(status_line("At the default threshold of 0.50" in full_text, "Default 0.50 threshold remains labelled"))

    # Figure legends and scripts.
    fig2_script = (PROJECT_ROOT / "scripts/submission_enhancement_v2/02_generate_discrimination_calibration_v2.py").read_text(encoding="utf-8")
    fig4_script = (PROJECT_ROOT / "scripts/submission_enhancement_v2/03_generate_paired_model_comparison_v2.py").read_text(encoding="utf-8")
    checks.append(status_line("fig.legend" in fig2_script and "ax.legend" not in fig2_script, "Figure 2 uses shared legend rather than panel legends over curves", critical=False))
    checks.append(status_line("Calculation" not in full_text or "0.05–0.80" in full_text or "0.05 to 0.80" in full_text, "Figure 3 calculation range stated"))
    checks.append(status_line("0.05–0.50" in full_text or "0.05 to 0.50" in full_text, "Figure 3 displayed range stated"))
    checks.append(status_line("Brier-score advantage" in fig4_script and "Brier-score advantage" in full_text, "Figure 4 uses Brier-score advantage label"))

    figure_stems = [
        OUT_ROOT / "main_figures/Figure2_discrimination_calibration_id_isolated",
        OUT_ROOT / "main_figures/Figure3_decision_curve_id_isolated_wave3",
        OUT_ROOT / "main_figures/Figure4_model_comparison_robustness",
    ]
    for stem in figure_stems:
        ok, detail = pdf_check(stem.with_suffix(".pdf"))
        checks.append(status_line(ok, f"{stem.relative_to(PROJECT_ROOT)}.pdf: {detail}"))
        ok, detail = png_check(stem.with_suffix(".png"))
        checks.append(status_line(ok, f"{stem.relative_to(PROJECT_ROOT)}.png: {detail}"))

    # Manuscript format/placeholders.
    checks.append(status_line(len(doc.inline_shapes) == 0, "No figures embedded in v9 DOCX"))
    checks.append(status_line("Table 1." in full_text and "Table 2." in full_text and "Table 3." in full_text, "Table numbering remains 1-3"))
    checks.append(status_line(all(f"Figure {i}." in full_text for i in [1, 2, 3, 4]), "Figure numbering remains 1-4"))
    checks.append(status_line("Additional file 1." in full_text and "Additional file 2." in full_text, "Additional-file numbering remains consistent"))
    checks.append(status_line("[AUTHOR ACTION" in full_text, "Author-action placeholders remain"))
    checks.append(status_line("Shanghai University of Sport" in full_text and "ethics committee" in lower, "Ethics placeholder remains"))
    checks.append(status_line("insert final public repository URL" in full_text, "Code repository placeholder remains"))

    render_ok, render_detail = render_v9()
    checks.append(status_line(render_ok, f"v9 DOCX render attempt: {render_detail}", critical=False))

    critical_failed = any(status == "FAIL" and critical for status, _message, critical in checks)
    warnings = any(status == "WARNING" for status, _message, _critical in checks)
    overall = "FAIL" if critical_failed else ("PASS WITH WARNINGS" if warnings else "PASS")

    lines = ["# v9 Final Consistency Check", "", f"Overall status: **{overall}**", "", "| Status | Check | Critical |", "|---|---|---:|"]
    for status, message, critical in checks:
        lines.append(f"| {status} | {message} | {critical} |")
    (LOG_DIR / "v9_final_consistency_check.md").write_text("\n".join(lines) + "\n", encoding="utf-8")

    manifest = [
        "# v9 Revision Manifest",
        "",
        f"- Checker status: {overall}",
        f"- Source manuscript: `{V8.relative_to(PROJECT_ROOT)}`",
        f"- Output manuscript: `{V9.relative_to(PROJECT_ROOT)}`",
        "- v8 was not overwritten.",
        "- No models were retrained or refit.",
        "- No raw data were modified.",
        "- No primary or strict-label outcome definitions were changed.",
        "- Figure 2 and Figure 4 were regenerated from existing fixed predictions/source tables.",
        "- Table 2 was rebuilt as an editable Word table with model-specific ID-isolated Wave 3 results only.",
        "- Wave 1 OOF rows were moved to `output/submission_enhancement_v2/tables/SuppTable_wave1_oof_model_performance.csv`.",
        f"- Render status: {'success' if render_ok else 'warning'}; {render_detail}",
        "",
        "## Files created or updated in this v9 pass",
        "- `output/submission_enhancement_v2/logs/v9_consistency_data_packet.md`",
        "- `output/submission_enhancement_v2/tables/Table2_main_model_validation_concise.csv`",
        "- `output/submission_enhancement_v2/tables/SuppTable_wave1_oof_model_performance.csv`",
        "- `output/submission_enhancement_v2/main_figures/Figure2_discrimination_calibration_id_isolated.pdf`",
        "- `output/submission_enhancement_v2/main_figures/Figure2_discrimination_calibration_id_isolated.png`",
        "- `output/submission_enhancement_v2/main_figures/Figure4_model_comparison_robustness.pdf`",
        "- `output/submission_enhancement_v2/main_figures/Figure4_model_comparison_robustness.png`",
        "- `manuscript/BMC_Geriatrics_possible_sarcopenia_revised_v9_final_consistency_corrected.docx`",
        "- `scripts/submission_enhancement_v2/02_generate_discrimination_calibration_v2.py`",
        "- `scripts/submission_enhancement_v2/03_generate_paired_model_comparison_v2.py`",
        "- `scripts/submission_enhancement_v2/09_check_v9_final_consistency.py`",
        "- `scripts/submission_enhancement_v2/10_integrate_manuscript_v9.py`",
    ]
    (MANIFEST_DIR / "v9_revision_manifest.md").write_text("\n".join(manifest) + "\n", encoding="utf-8")

    print(overall)
    if critical_failed:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
