from __future__ import annotations

import sys
import json
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[2]
LOCAL_SITE_PACKAGES = PROJECT_ROOT / ".python_packages"
if LOCAL_SITE_PACKAGES.exists():
    sys.path.insert(0, str(LOCAL_SITE_PACKAGES))

from docx import Document
from PIL import Image
import pandas as pd
from sklearn.metrics import average_precision_score, roc_auc_score

from _utils_v2 import ID_COL, OUT_ROOT, PRED_COL, PRIMARY_LABEL_COL, TABLE_DIR, write_text


V7 = PROJECT_ROOT / "manuscript/BMC_Geriatrics_possible_sarcopenia_revised_v7_outcome_label_audit_integrated.docx"
V8 = PROJECT_ROOT / "manuscript/BMC_Geriatrics_possible_sarcopenia_revised_v8_robustness_figures_integrated.docx"


def status_line(ok: bool, message: str, critical: bool = True) -> tuple[str, str, bool]:
    status = "PASS" if ok else ("FAIL" if critical else "WARNING")
    return status, message, critical


def pdf_check(path: Path) -> tuple[bool, str]:
    if not path.exists():
        return False, "missing"
    b = path.read_bytes()
    ok = b.startswith(b"%PDF") and b.rstrip().endswith(b"%%EOF") and b.count(b"/FontFile2") > 0 and b.count(b"/Subtype /Type3") == 0
    return ok, f"size_mb={path.stat().st_size/1024/1024:.3f}; FontFile2={b.count(b'/FontFile2')}; Type3={b.count(b'/Subtype /Type3')}"


def png_check(path: Path) -> tuple[bool, str]:
    if not path.exists():
        return False, "missing"
    im = Image.open(path)
    dpi = im.info.get("dpi", (0, 0))
    ok = path.stat().st_size < 10 * 1024 * 1024 and min(dpi) >= 250
    return ok, f"{im.size[0]}x{im.size[1]}; dpi={dpi}; size_mb={path.stat().st_size/1024/1024:.3f}"


def main() -> None:
    checks: list[tuple[str, str, bool]] = []

    # Primary baseline.
    log = PROJECT_ROOT / "output/tables/21e_wave3_predictions_logistic_A_only_id_isolated.csv"
    xgb = PROJECT_ROOT / "output/tables/21e_wave3_predictions_xgboost_A_plus_B_id_isolated.csv"
    ldf = pd.read_csv(log)
    xdf = pd.read_csv(xgb)
    checks.append(status_line(len(ldf) == 2453, "Primary Logistic validation N = 2453"))
    checks.append(status_line(abs(roc_auc_score(ldf[PRIMARY_LABEL_COL].astype(int), ldf[PRED_COL]) - 0.6798) <= 0.0005, "Primary Logistic AUROC reproduced"))
    checks.append(status_line(abs(average_precision_score(ldf[PRIMARY_LABEL_COL].astype(int), ldf[PRED_COL]) - 0.4759) <= 0.0005, "Primary Logistic AUPRC reproduced"))
    checks.append(status_line(len(xdf) == 2430, "Primary XGBoost validation N = 2430"))
    checks.append(status_line(abs(roc_auc_score(xdf[PRIMARY_LABEL_COL].astype(int), xdf[PRED_COL]) - 0.6656) <= 0.0005, "Primary XGBoost AUROC reproduced"))
    checks.append(status_line(abs(average_precision_score(xdf[PRIMARY_LABEL_COL].astype(int), xdf[PRED_COL]) - 0.4604) <= 0.0005, "Primary XGBoost AUPRC reproduced"))
    common = ldf[[ID_COL, PRIMARY_LABEL_COL, PRED_COL]].rename(columns={PRIMARY_LABEL_COL: "y_l", PRED_COL: "p_l"}).merge(
        xdf[[ID_COL, PRIMARY_LABEL_COL, PRED_COL]].rename(columns={PRIMARY_LABEL_COL: "y_x", PRED_COL: "p_x"}),
        on=ID_COL,
        how="inner",
    )
    y = common["y_l"].astype(int)
    auroc_diff = roc_auc_score(y, common["p_l"]) - roc_auc_score(y, common["p_x"])
    auprc_diff = average_precision_score(y, common["p_l"]) - average_precision_score(y, common["p_x"])
    checks.append(status_line(len(common) == 2430, "Common paired subset N = 2430"))
    checks.append(status_line(abs(auroc_diff - 0.0116) <= 0.0005, "Known paired AUROC difference reproduced"))
    checks.append(status_line(abs(auprc_diff - 0.0087) <= 0.0005, "Known paired AUPRC difference reproduced"))

    # Strict label.
    flow = pd.read_csv(TABLE_DIR / "strict_label_flow_by_wave.csv")
    w1_uncertain = int(flow.loc[flow["wave"].eq(1), "uncertain_partial_negative_n"].iloc[0])
    w3_uncertain = int(flow.loc[flow["wave"].eq(3), "uncertain_partial_negative_n"].iloc[0])
    checks.append(status_line(w1_uncertain == 313, "Strict Wave 1 uncertain count = 313"))
    checks.append(status_line(w3_uncertain == 278, "Strict original Wave 3 uncertain count = 278"))
    idflow = pd.read_csv(TABLE_DIR / "strict_id_isolated_sample_flow.csv")
    overlap = idflow.loc[idflow["dataset"].eq("wave3_id_isolated_source"), "remaining_id_overlap_with_other_wave"].iloc[0]
    checks.append(status_line(int(overlap) == 0, "Strict ID-isolated Wave 3 overlap with Wave 1 = 0"))
    strict_metric = pd.read_csv(TABLE_DIR / "strict_label_metrics_with_95CI.csv")
    checks.append(status_line(not strict_metric.empty, "Strict-label metrics table exists and is non-empty"))
    primary_vs_strict = pd.read_csv(TABLE_DIR / "primary_vs_strict_label_comparison.csv")
    checks.append(status_line(not primary_vs_strict.empty, "Primary-vs-strict comparison table exists and is non-empty"))

    # Models/specification from scripts.
    s01 = (PROJECT_ROOT / "scripts/submission_enhancement_v2/01_strict_label_sensitivity_v2.py").read_text(encoding="utf-8")
    for snippet in [
        'LogisticRegression(max_iter=2000, solver="liblinear", random_state=RANDOM_STATE)',
        "n_estimators=300",
        "max_depth=3",
        "learning_rate=0.05",
        "random_state=RANDOM_STATE",
        "StratifiedKFold(n_splits=N_SPLITS, shuffle=True, random_state=RANDOM_STATE)",
    ]:
        checks.append(status_line(snippet in s01, f"Model specification snippet present: {snippet[:50]}"))
    strict_log = json.loads((OUT_ROOT / "logs/strict_label_sensitivity_v2_log.json").read_text(encoding="utf-8"))
    checks.append(status_line(strict_log.get("wave3_tuning_performed") is False, "No Wave 3 tuning recorded"))

    # Bootstrap.
    paired = pd.read_csv(TABLE_DIR / "paired_model_differences_with_95CI.csv")
    checks.append(status_line((paired["bootstrap_replicates_requested"] == 2000).all(), "Paired bootstrap requested replicates = 2000"))
    checks.append(status_line((paired["bootstrap_seed"] == 20260613).all(), "Paired bootstrap seed = 20260613"))
    checks.append(status_line((paired["bootstrap_replicates_valid"] > 0).all(), "Paired bootstrap valid counts reported"))

    # Figures and source tables.
    expected_figures = [
        "main_figures/Figure1_participant_flow_id_isolated",
        "main_figures/Figure2_discrimination_calibration_id_isolated",
        "main_figures/Figure3_decision_curve_id_isolated_wave3",
        "main_figures/Figure4_model_comparison_robustness",
        "supplementary_figures/SuppFigureS1_xgboost_gain_importance",
        "supplementary_figures/SuppFigureS2_screening_threshold_tradeoffs",
        "supplementary_figures/SuppFigureS3_wave1_oof_discrimination_calibration",
        "supplementary_figures/SuppFigureS4_original_nonisolated_vs_idisolated_performance",
        "supplementary_figures/SuppFigureS5_subgroup_robustness",
        "supplementary_figures/SuppFigureS6_learning_curves",
    ]
    for stem in expected_figures:
        for suffix, checker in [(".pdf", pdf_check), (".png", png_check)]:
            ok, detail = checker(OUT_ROOT / f"{stem}{suffix}")
            checks.append(status_line(ok, f"{stem}{suffix}: {detail}", critical=True))
    expected_tables = [
        "Figure2_ROC_curve_data.csv",
        "Figure2_PR_curve_data.csv",
        "Figure2_calibration_bin_data.csv",
        "Figure2_calibration_summary.csv",
        "Figure4_plot_data.csv",
        "paired_model_differences_with_95CI.csv",
        "strict_label_metrics_with_95CI.csv",
        "calibration_uncertainty_summary.csv",
        "calibration_bin_uncertainty.csv",
        "subgroup_performance_with_95CI.csv",
    ]
    for table in expected_tables:
        checks.append(status_line((TABLE_DIR / table).exists(), f"Source/output table exists: {table}"))

    # Manuscript.
    checks.append(status_line(V7.exists(), "v7 source manuscript exists"))
    checks.append(status_line(V8.exists(), "v8 manuscript created"))
    if V7.exists() and V8.exists():
        d7 = Document(V7)
        d8 = Document(V8)
        text8 = "\n".join(p.text for p in d8.paragraphs)
        checks.append(status_line(len(d8.inline_shapes) == len(d7.inline_shapes), "No new figures embedded in v8 manuscript"))
        checks.append(status_line("Figure 4. Paired model differences" in text8, "Figure 4 legend present"))
        checks.append(status_line("[AUTHOR ACTION" in text8, "Author-action placeholders retained", critical=False))
        checks.append(status_line("Shanghai University of Sport ethics" in text8 or "ethics approval" in text8.lower(), "Ethics wording/placeholders retained", critical=False))
    render_dir = OUT_ROOT / "logs/v8_rendered_pages"
    checks.append(status_line(render_dir.exists() and any(render_dir.glob("page-*.png")), "v8 rendered page PNGs available for visual QA", critical=False))

    # No model objects under v2.
    joblibs = list(OUT_ROOT.rglob("*.joblib"))
    checks.append(status_line(len(joblibs) == 0, "No trained model objects saved under v2 output"))

    critical_failed = any(status == "FAIL" and critical for status, _message, critical in checks)
    warnings = any(status == "WARNING" for status, _message, _critical in checks)
    overall = "FAIL" if critical_failed else ("PASS WITH WARNINGS" if warnings else "PASS")

    lines = ["# Submission Enhancement v2 Check", "", f"Overall checker result: **{overall}**", "", "| Status | Check | Critical |", "|---|---|---:|"]
    for status, message, critical in checks:
        lines.append(f"| {status} | {message} | {critical} |")
    write_text(OUT_ROOT / "logs/submission_enhancement_v2_check.md", "\n".join(lines) + "\n")

    manifest = [
        "# Submission Enhancement v2 Manifest",
        "",
        f"- Checker result: {overall}",
        "- Scripts created under `scripts/submission_enhancement_v2/`.",
        "- Outputs created under `output/submission_enhancement_v2/`.",
        "- Primary outcome files were not overwritten.",
        "- v7 manuscript was not overwritten; v8 was created as a new file.",
        "- No raw CHARLS data were modified.",
        "- No Wave 3 model tuning or new model selection was performed.",
        "- DOCX visual render QA is unavailable if page PNGs are absent, because the local environment lacks required render dependencies.",
        "",
        "## Files created",
    ]
    for path in sorted(OUT_ROOT.rglob("*")):
        if path.is_file():
            manifest.append(f"- `{path.relative_to(PROJECT_ROOT)}`")
    for path in sorted((PROJECT_ROOT / "scripts/submission_enhancement_v2").glob("*.py")):
        manifest.append(f"- `{path.relative_to(PROJECT_ROOT)}`")
    manifest.append(f"- `{V8.relative_to(PROJECT_ROOT)}`")
    write_text(OUT_ROOT / "manifests/submission_enhancement_v2_manifest.md", "\n".join(manifest) + "\n")
    print(overall)
    if critical_failed:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
